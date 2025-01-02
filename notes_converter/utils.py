"""A module containing all the functions that make notes_converter work.
"""

import csv
import functools
import getpass
import json
import re
import sqlite3
from collections import namedtuple
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Sequence, Union

import psutil
import pytz
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.opc.oxml import qn
from docx.oxml import OxmlElement

ROOT_PATH = Path.home()
CWD = Path.cwd()

TEST_DATA_PATH = CWD / "tests"

DATA_PATH = CWD / "data"
TEMPLATE_PATH = CWD / "templates"

FIELD_NAMES = [
    "type",
    "title",
    "note_text",
    "source_location",
    "tags",
    "notebooks",
    "study_set",
    "last_updated",
    "created",
    "highlight",
]

TAGS = set()
NOTEBOOKS = set()


def report_error(func):
    @functools.wraps(func)
    def wrapper_report_error(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except TypeError as e:
            print(f"TypeError: {e}")
            for item in args[0]:
                try:
                    func([item])
                except TypeError:
                    print(f"Problematic item: {item.title}")
                    with open(
                        "Problematic_items_titles.txt", "a", encoding="utf-8"
                    ) as f:
                        f.write(item.title + "\n")
            raise

    return wrapper_report_error


def remove_duplicates(func):
    """Remove duplicate notes."""

    # TODO: Add more tailored duplicate removal.
    # For example, if a note is the same but the tags or notebooks are
    # different. In these cases, the note most recently updated should
    # be kept and the other discarded.

    @functools.wraps(func)
    def wrapper_remove_duplicates(*args, **kwargs):
        removed_exact_duplicates = []
        for note in func(*args, **kwargs):
            if note not in removed_exact_duplicates:
                removed_exact_duplicates.append(note)
        return removed_exact_duplicates

    return wrapper_remove_duplicates


class NoAvailableTemplate(Exception):
    pass


def load_json(path):
    """Load a `.json` file in `utf-8` encoding.

    Parameters
    ----------
    path : A string or Path object.

    Returns
    -------
    The contents of the `.json` file.
    """
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_csv_as_dict(
    path: Union[Path, str],
    field_names: List[str],
) -> csv.DictReader:
    """Load a CSV into a `csv.DictReader` iterable.

    Parameters
    ----------
    path : A string or Path object to the file.
    field_names : A list of strings mapping to the csv fields.

    Returns
    -------
    A `csv.DictReader` iterable.
    """
    return csv.DictReader(open(path, encoding="utf-8"), fieldnames=field_names)


@remove_duplicates
def load_csv_files(files: Sequence[Union[Path, str]], field_names: List[str]):
    """Load multiple `csv` files and return a merged list.

    Parameters
    ----------
    files : A list containing the paths to the files.
        These can be either `str` or `Path` objects.
    field_names : A list of strings
        Containing the fields to which to map the csv columns.

    Returns
    -------
    A merged list of `dict`s.
    """
    notes = []
    for file in files:
        reader = load_csv_as_dict(file, field_names=field_names)
        next(reader)  # Skip titles (first line)
        notes += [note for note in reader]
    return notes


def open_database(database):
    pass


def commit_to_database(
    database, files: Sequence[Union[Path, str]], field_names: List[str]
):
    for file in files:
        reader = load_csv_as_dict(file, field_names=field_names)
        next(reader)  # Skip titles (first line)
        with sqlite3.connect(database) as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT into VALUES(?, ?, ?)")


class SystemMemory:
    def __init__(self) -> None:
        self._memory_info = psutil.virtual_memory()
        self._current_memory = self._memory_info.available
        self._current_storage = self._memory_info.total

    def check_memory(self, megabytes) -> bool:
        """Check the available system memory (RAM) against `megabytes`.

        Parameters
        ----------
        - megabytes : An estimated amount of required memory in megabytes.

        Returns
        -------
        A boolean value: `True` or `False`.
        """
        # Convert to megabytes:
        needed_memory = megabytes
        current_memory = self._current_memory / (1024 * 1024)

        return True if needed_memory < current_memory else False

    def check_storage(self, megabytes) -> bool:
        """Check the available system storage (disk space) against
        `megabytes`.

        Parameters
        ----------
        - megabytes : An estimated amount of required storage
            in megabytes.

        Returns
        -------
        A boolean value: `True` or `False`.
        """
        # Convert to megabytes:
        needed_storage = megabytes
        current_storage = self._current_storage / (1024 * 1024)
        return True if needed_storage < current_storage else False


def check_file_size(paths):
    """Tally the size of all provided files, in bytes, and
    return the sum in megabytes."""
    return sum([Path(path).stat().st_size / (1024 * 1024) for path in paths])


def build_notes(
    notes_list: List[Dict[str, Union[str, List[str]]]], field_names: List[str]
):
    """Build a list of notes using a `namedtuple` object.

    Parameters
    ----------
    notes_list : A `list` of dictionaries.

    Returns
    -------
    A list of `namedtuple` objects.
    """

    Note = namedtuple("Note", field_names=field_names)

    notes = []
    for n in notes_list:
        n = process_note(n)
        values = [value if value else "" for value in n.values()]
        note = Note(*values)
        notes.append(note)
    return notes


def sort_notes_by_title_and_verse(notes, title_order):
    """Sort notes by title and verse.

    Parameters
    ----------
    notes : A list of Note objects with `.title`, `.date` and `.body`
        attributes.

    title_order : A list of strings in the desired order for the output notes.

    Returns
    -------
    A list of `Note` objects sorted by title and verse.
    """
    indexed_titles = {title: index for index, title in enumerate(title_order)}
    chapter_verse_pattern = re.compile(r"\W(\d+|\d+:\d+)|;|:.*")

    def extract_chapter_and_verse(title):
        _match = [i for i in chapter_verse_pattern.findall(title) if i.strip()]
        index = 1
        if _match:
            chapter = int(_match[0])
            verse = int(_match[index]) if index < len(_match) else 0
            return chapter, verse
        return 0, 0

    def arrange_by_title_and_verse(note):
        title_without_verse = chapter_verse_pattern.sub("", note.title)
        chapter_index = indexed_titles.get(title_without_verse, float("inf"))
        chapter, verse = extract_chapter_and_verse(note.title)
        return chapter_index, chapter, verse

    return sorted(notes, key=arrange_by_title_and_verse)


def process_note(note):
    """Convert a note's `note_text`, `tags` and `notebooks` to lists,
    and remove the date and ruler, and update the `TAGS` and `NOTEBOOKS`
    constants.
    """
    n = convert_note_text_to_list(note)
    _n = convert_note_identifiers_to_list(n)
    cleaned_n = remove_headers(_n)

    add_tags(TAGS, cleaned_n["tags"])
    add_tags(NOTEBOOKS, cleaned_n["notebooks"])

    return cleaned_n


# process_note() helper functions


def add_tags(tags_set, new_tags) -> None:
    tags_set.update(new_tags)


def strip_artifacts(line: str) -> Union[str, None]:
    """Strip unwanted parts of a list."""
    if line.startswith("[") or line.startswith("-----"):
        return ""
    return line


def remove_headers(note):
    """Remove unwanted text at the head of a list."""
    n = [strip_artifacts(_) for _ in note["note_text"]]
    cleaned_n = [_ for _ in n if _.strip()]  # Remove empty values
    note["note_text"] = cleaned_n
    return note


def convert_note_text_to_list(note):
    """Convert `note`'s `note_text` value to a list using `re.findall()`.

    Parameters
    ----------
    note : A `dict` object.

    Returns
    -------
    A `dict` object with only its `["note_text"]` value changed.
    """
    paragraphs = re.compile(r"([^\n]+(?:\n(?!\n)[^\n]+)*)")
    p = paragraphs.findall(note["note_text"])
    text = [i.replace("\n", " ") for i in p]
    note["note_text"] = text
    return note


def convert_note_identifiers_to_list(note: Dict) -> Dict:
    """Convert a note's `["tags"]` and `["notebooks"]` values to lists.

    Parameters
    ----------
    note : A `dict` object.

    Returns
    -------
    A `dict` object with only its `["tags"]` and `["notebooks"]` values changed.
    """
    tags: List[str] = note["tags"].split("; ")
    note["tags"] = tags
    notebooks: List[str] = note["notebooks"].split("; ")
    note["notebooks"] = notebooks
    return note


# Other


def build_source_names(url: str):
    hyperlink_name = r"""
        https:\/\/www\.churchofjesuschrist\.org\/study\/scriptures\/ # Match
        # the static part of the url
        ([a-z]{2,5})\/  # Capture the two or three-letter book identifier
        ([a-z]+)\/  # Capture the book name
        (\d+)\/  # Capture the chapter number
        \?lang=[a-z]{3}&id=p(\d+)  # Match the query parameters and capture
        # the paragraph number.
    """
    names = re.compile(hyperlink_name, re.VERBOSE)
    return names.findall(url)


def write_to_txt(notes, output_path):
    """Write the given notes to a `.txt` file."""
    with open(output_path, "w", encoding="utf-8") as f:
        for note in notes:
            f.write("\n\n" + note.title + "\n")
            f.write(convert_datetime(note.created) + "\n\n")
            for body in note.note_text:
                f.write(body + " ")
            f.write("\n" + note.source_location)


def write_to_docx(
    notes,
    output_path: Union[str, Path],
    template_path: Union[str, Path, None],
):
    """Write notes to a styled Word document.

    Parameters
    ----------
    notes : A list of `note` objects.
    output_path : The location to save the Word document.
    template_path : A path to a Word document template.
        `None` means that the default template will be used.

    Returns
    -------
    A styled Word document in the given `output_path`.
    """
    # Is there a custom template path?
    if template_path:
        template = Path(template_path)
        valid_template = template.exists() and template.suffix == ".docx"
        if not valid_template:
            template = TEMPLATE_PATH / "default.docx"
    else:
        template = TEMPLATE_PATH / "default.docx"

    try:
        doc = Document(str(template))
    except PackageNotFoundError:
        # This error is raised only when default.docx is open in another
        # application. Aside from development, that should (in theory)
        # never happen in production.
        raise NoAvailableTemplate

    file_name = Path(output_path).stem

    # Clear existing template data
    doc._body.clear_content()
    doc.add_heading(file_name, level=0)

    for note in notes:
        doc.add_heading(note.title, level=1)
        doc.add_paragraph(convert_datetime(note.created), style="Date")

        for value, body in enumerate(note.note_text):
            if value == 0:  # Allow for no text indent on first paragraph.
                doc.add_paragraph(body, style="Head")
                continue
            doc.add_paragraph(body, style="Normal")

        # TODO: an add_run() will be needed to use Word's
        # built-in Hyperlink style. Can add_hyperlink() be
        # modified to use add_run() and to add the hyperlink
        # to that?
        p = doc.add_paragraph(style="Link")
        add_hyperlink(p, note.source_location, "Source: ")

    # Add document properties
    doc.core_properties.author = getpass.getuser()
    doc.core_properties.comments = "Document generated by a script."

    doc.save(str(output_path))


# write_to_docx() helper functions


def convert_datetime(note_time: str) -> str:
    """Convert the time to a human-readable format."""
    default_time = datetime.fromisoformat(note_time.replace("Z", "+00:00"))
    dt = default_time.replace(tzinfo=pytz.utc)
    pacific_tz = pytz.timezone("America/Los_Angeles")
    dt_pacific = dt.astimezone(pacific_tz)
    return dt_pacific.strftime("%B %d, %Y, %I:%M %p %Z")


def add_hyperlink(paragraph, url, text, color="#0000EE", underline=None):
    """Place a hyperlink within a `paragraph` object.

    Parameters
    ----------
    paragraph : The `paragraph` object to which to add the hyperlink.
    url : A string containing a given url.
    text : The text displayed for the url.
    color : The color of the hyperlink.
    underline : Whether an underline is applied.

    Returns
    -------
    A `hyperlink` object with the hyperlink
    """
    # Get access to the document.xml.rels file and get a new relation id value:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    # Create the w:hyperlink tag and add needed values:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element:
    new_run = OxmlElement("w:r")

    # Create a w:rPr element:
    rPr = OxmlElement("w:rPr")

    # Add provided color (if any):
    if color:
        c = OxmlElement("w:color")
        c.set(qn("ct:val"), color)
        rPr.append(c)

    # Remove underline if requested:
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("ct:val"), "single")
        rPr.append(u)
    else:
        u = OxmlElement("w:u")
        u.set(qn("ct:val"), "none")
        rPr.append(u)

    # Join all xml elements and add hyperlink text to w:r element:
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
